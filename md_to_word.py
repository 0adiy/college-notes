import re
import os
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml import OxmlElement
import requests

ROLL_NO = "A02"
INPUT_MARKDOWN_FILE = "./Soft Computing/all labs.md"


def set_heading_style(doc, level, font_size, font_color):
    style_name = f"Heading {level}"
    style = doc.styles[style_name]

    if not style:
        # Create a new style if it does not exist
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    # Set font size
    font = style.font
    font.size = Pt(font_size)

    # Set font color
    font.color.rgb = RGBColor(*font_color)

    return style


def add_paragraph(doc, text, style=None, font_color=None):
    p = doc.add_paragraph(text, style=style)
    if font_color:
        run = p.runs[0]
        run.font.color.rgb = RGBColor(*font_color)
    return p


def add_heading(doc, text, level, font_size, font_color):
    set_heading_style(doc, level, font_size, font_color)
    doc.add_heading(text, level=level)


def is_code_block(paragraph_text):
    # Check if a paragraph starts and ends with triple backticks
    return paragraph_text.startswith("```") and paragraph_text.endswith("```")


def clean_code_text(code_text):
    # Regular expression to match lines with triple backticks
    print(code_text)
    cleaned_code_text = re.sub(r"^```.*$", "", code_text, flags=re.MULTILINE)
    print(cleaned_code_text)
    return cleaned_code_text.strip()


def add_code_block(doc, code_text):
    p = doc.add_paragraph(style="Normal")  # Use default paragraph style
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10.5)
    p.style.font.name = "Normal"
    p.style.font.size = Pt(10.5)
    p.style.font.color.rgb = RGBColor(0, 0, 0)  # Default color for code


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_header(doc, header_text):
    section = doc.sections[0]
    header = section.header
    p = header.add_paragraph()
    p.text = header_text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create and add a page number field to the paragraph
    page_number_field = OxmlElement("w:fldSimple")
    page_number_field.set(qn("w:instr"), "PAGE")
    page_number_field.set(qn("w:result"), "1")

    run = p.add_run()
    run._element.append(page_number_field)


# Your existing functions...


def add_image(doc, image_path):
    # Add an image to the document.
    try:
        doc.add_picture(image_path)
    except Exception as e:
        print(f"Failed to add image {image_path}: {e}")


def markdown_to_word(md_text, output_file):
    # Convert Markdown to HTML
    html = markdown.markdown(md_text, extensions=["fenced_code"])
    with open("md.html", "w", encoding="utf-8") as f:
        f.write(html)

    # Create a new Document
    doc = Document()

    # Add header and footer
    add_header(doc, ROLL_NO)
    add_footer(doc)

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    first_heading = True
    for element in soup:
        if element.name == "h1":  # type: ignore
            if not first_heading:
                add_page_break(doc)
            add_heading(
                doc, element.get_text(), level=1, font_size=22, font_color=(0, 0, 0)
            )
            first_heading = False
        elif element.name == "h2":
            add_heading(
                doc, element.get_text(), level=2, font_size=18, font_color=(0, 0, 0)
            )
        elif element.name == "p":
            # Process paragraph text and images
            for sub_element in element:
                if sub_element.name == "img":
                    img_url = sub_element.get("src")
                    if img_url:
                        # Download the image if it's a URL
                        if img_url.startswith("http"):
                            img_data = requests.get(img_url).content
                            img_name = os.path.basename(img_url)
                            with open(img_name, "wb") as handler:
                                handler.write(img_data)
                            add_image(doc, img_name)
                            os.remove(img_name)  # Clean up downloaded image
                        else:
                            # If it's a local file, just use its path
                            add_image(doc, img_url)
            else:
                if element.get_text() != "":
                    add_paragraph(doc, element.get_text(), font_color=(0, 0, 0))
        elif element.name == "pre":
            code_text = element.get_text()
            add_code_block(doc, code_text)
        elif element.name == "img":  # Detect image tags
            img_url = element.get("src")
            if img_url:
                # Download the image if it's a URL
                if img_url.startswith("http"):
                    img_data = requests.get(img_url).content
                    img_name = os.path.basename(img_url)
                    with open(img_name, "wb") as handler:
                        handler.write(img_data)
                    add_image(doc, img_name)
                    os.remove(img_name)  # Clean up downloaded image
                else:
                    # If it's a local file, just use its path
                    add_image(doc, img_url)

    # Save the document
    doc.save(output_file)


if __name__ == "__main__":
    input_markdown_file = INPUT_MARKDOWN_FILE
    output_word_file = "output.docx"  # Desired output Word file

    with open(input_markdown_file, "r", encoding="utf-8") as file:
        md_text = file.read()

    markdown_to_word(md_text, output_word_file)
